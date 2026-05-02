#!/usr/bin/env python3
"""生成使用手册和开发手册 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ========== 通用工具函数 ==========

FONT_TITLE = '方正小标宋简体'
FONT_HEADING = '黑体'
FONT_SUBHEADING = '楷体_GB2312'
FONT_BODY = '仿宋_GB2312'
FONT_CODE = 'Courier New'

def set_cell_font(cell, text, font_name=FONT_BODY, size=Pt(12), bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, FONT_HEADING, Pt(11), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # 表头背景色
        shading = table.rows[0].cells[i]._element
        tcPr = shading.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = shading.makeelement(qn('w:tcPr'), {})
            shading.insert(0, tcPr)
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = tcPr.makeelement(qn('w:shd'), {})
            tcPr.append(shd)
        shd.set(qn('w:fill'), '4A5568')
    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            set_cell_font(table.rows[ri+1].cells[ci], str(cell_text), FONT_BODY, Pt(10))
    # 列宽
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, width in enumerate(col_widths):
                row.cells[ci].width = Cm(width)
    doc.add_paragraph()
    return table

def add_body(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(12)
    run.font.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = Pt(16)
    run.font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_SUBHEADING
    run.font.size = Pt(14)
    run.font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_SUBHEADING)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_info_box(doc, title, text):
    """添加信息提示框（用表格模拟）"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(f'【{title}】{text}')
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    # 浅蓝背景
    shading = cell._element
    tcPr = shading.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = shading.makeelement(qn('w:tcPr'), {})
        shading.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = tcPr.makeelement(qn('w:shd'), {})
        tcPr.append(shd)
    shd.set(qn('w:fill'), 'EBF5FF')
    doc.add_paragraph()

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(12)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    if level > 0:
        p.paragraph_format.left_indent = Cm(level * 0.74)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = FONT_CODE
    run.font.size = Pt(9)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    # 浅灰背景
    shd = p._element.makeelement(qn('w:shd'), {})
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shd)
    return p

# ========== 使用手册 ==========

def generate_user_manual():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('开放智慧助教')
    run.font.name = FONT_TITLE
    run.font.size = Pt(36)
    run.font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_TITLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('使 用 手 册')
    run.font.name = FONT_TITLE
    run.font.size = Pt(30)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_TITLE)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('基于教师个人知识库的智能语音交互教学系统')
    run.font.name = FONT_BODY
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    doc.add_page_break()

    # ===== 第一章 系统简介 =====
    add_heading1(doc, '第一章 系统简介')

    add_heading2(doc, '1.1 系统概述')
    add_body(doc, '开放智慧助教（Open Teacher）是一款基于教师个人知识库的智能语音交互教学系统，专为非全日制成人学生设计，解决工学矛盾导致的学生无法与教师常态化交流的问题。系统支持教师创建专属数字助教，将个人教学经验和知识体系数字化，实现7×24小时个性化语音辅导，并通过学生记忆系统实现"越教越懂学生"的渐进式教学。', indent=True)

    add_heading2(doc, '1.2 核心特色')
    add_bullet(doc, '语音实时交互：基于Web Audio API的实时语音对话，支持自动语音检测、回声防护、手动打断，体验接近真人通话')
    add_bullet(doc, '价值观领域约束：AI助教严格限定在教师专业范围内回答，拒绝回答与教学无关的问题，确保教育安全')
    add_bullet(doc, '学生记忆系统：自动追踪知识掌握程度（已掌握/学习中/薄弱三档），每次对话都基于历史记忆个性化教学')
    add_bullet(doc, '教师专属知识库：每位教师创建独立知识库，AI回答基于教师上传的专业资料，确保内容权威性')
    add_bullet(doc, '数字人形象与音频波纹：使用教师真实照片作为数字人形象，语音播报时展示音频波纹动画，直观反馈播放状态')
    add_bullet(doc, '辅助图示生成：对抽象概念自动生成辅助图示，帮助学生直观理解')

    add_heading2(doc, '1.3 用户角色')
    add_table(doc,
        ['角色', '说明', '主要操作'],
        [
            ['管理员', '系统最高权限', '创建/管理教师账号、管理学生账号、启用/禁用账号'],
            ['教师', '知识库构建者', '设置助教档案、上传知识库、管理数字人、查看学生记忆'],
            ['学生', '学习使用者', '选择助教、语音/文字交互、查看学习记忆'],
        ],
        col_widths=[3, 4, 9]
    )

    # ===== 第二章 快速开始 =====
    add_heading1(doc, '第二章 快速开始')

    add_heading2(doc, '2.1 系统访问')
    add_body(doc, '通过浏览器访问系统部署地址即可使用，支持电脑和手机浏览器，无需安装任何软件。', indent=True)

    add_heading2(doc, '2.2 使用流程概览')
    add_table(doc,
        ['步骤', '操作者', '操作内容'],
        [
            ['1', '管理员', '登录后台（默认账号admin/admin123），创建教师账号'],
            ['2', '教师', '使用分配的账号登录，完成4步设置向导'],
            ['3', '学生', '输入姓名，选择已配置好的助教'],
            ['4', '学生', '点击电话按钮开始语音交互，或输入文字对话'],
        ],
        col_widths=[2, 3, 11]
    )

    # ===== 第三章 源码部署 =====
    add_heading1(doc, '第三章 源码部署')

    add_info_box(doc, '适用场景', '本章适用于使用项目源代码自行部署的情况。如果使用已部署好的在线服务，可直接跳至第四章开始使用。')

    add_heading2(doc, '3.1 方式一：Coze 编程平台部署（推荐）')
    add_body(doc, '使用 Coze 编程平台部署是最简便的方式，环境配置和依赖安装均可通过自然语言交流自动完成，无需手动配置。', indent=True)

    add_heading3(doc, '部署步骤')
    add_bullet(doc, '第1步：将项目源码压缩为 ZIP 文件')
    add_bullet(doc, '第2步：登录 Coze 编程平台（coze.cn），创建新项目')
    add_bullet(doc, '第3步：上传源码 ZIP 文件到项目中')
    add_bullet(doc, '第4步：通过自然语言对话告知平台需要配置的环境信息，例如：')
    add_code_block(doc, """请配置以下环境变量：
- 数据库：Supabase，URL为 https://xxxxx.supabase.co，
  anon key 为 eyJhbGciOi...，service_role key 为 eyJhbGciOi...
- 对象存储：S3 兼容存储，endpoint 为 https://oss-cn-beijing.aliyuncs.com，
  bucket 名称为 your-bucket-name""")
    add_bullet(doc, '第5步：平台自动完成环境配置、依赖安装和服务启动')
    add_bullet(doc, '第6步：访问平台分配的域名即可使用')

    add_info_box(doc, '优势', 'Coze 编程平台自动处理 AI 能力认证（LLM/ASR/TTS/知识库/图片生成）、依赖安装、端口配置、热更新等，无需手动操作。')

    add_heading3(doc, '需要提前准备的信息')
    add_body(doc, '即使是 Coze 平台部署，以下两个外部服务仍需自行注册并提供连接信息：', indent=True)
    add_table(doc,
        ['服务', '用途', '申请地址', '需要提供的信息'],
        [
            ['Supabase', '学生记忆系统数据库', 'supabase.com（有免费额度）', 'Project URL、anon key、service_role key'],
            ['S3 对象存储', '教师头像文件存储', '阿里云OSS/腾讯云COS/AWS S3', 'Endpoint URL、Bucket Name'],
        ],
        col_widths=[3, 4, 4, 5]
    )
    add_info_box(doc, '提示', 'Coze 平台的 AI 能力（LLM/ASR/TTS/知识库/图片生成）由平台自动提供认证，无需额外申请或配置。')

    add_heading2(doc, '3.2 方式二：独立服务器部署')
    add_body(doc, '如需在自有服务器上部署，需要手动完成所有环境配置。仅适用于有运维经验的用户。', indent=True)

    add_heading3(doc, '系统要求')
    add_table(doc,
        ['项目', '要求'],
        [
            ['操作系统', 'Linux / macOS / Windows（WSL2）'],
            ['Node.js', 'v20 及以上（推荐 v24）'],
            ['pnpm', 'v9.0.0 及以上'],
            ['内存', '最低 2GB，推荐 4GB'],
        ],
        col_widths=[4, 12]
    )

    add_heading3(doc, '需要申请的外部服务')
    add_body(doc, '独立部署需要额外申请 Coze 平台账号，以获取 AI 能力的访问认证：', indent=True)
    add_table(doc,
        ['服务', '用途', '说明'],
        [
            ['Coze 平台', 'LLM/ASR/TTS/知识库/图片生成', '需注册 coze.cn，通过 Coze CLI 认证调用'],
            ['Supabase', '学生记忆系统数据库', '注册 supabase.com，获取 URL 和密钥'],
            ['S3 对象存储', '教师头像存储', '选择 S3 兼容存储服务（OSS/COS/S3）'],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading3(doc, '环境变量配置')
    add_body(doc, '在项目根目录创建 .env 文件：', indent=True)
    add_code_block(doc, """# Supabase 数据库（必填）
COZE_SUPABASE_URL=https://xxxxx.supabase.co
COZE_SUPABASE_ANON_KEY=eyJhbGciOi...
COZE_SUPABASE_SERVICE_ROLE_KEY=eyJhbGciOi...

# S3 对象存储（必填）
COZE_BUCKET_ENDPOINT_URL=https://oss-cn-beijing.aliyuncs.com
COZE_BUCKET_NAME=your-bucket-name

# 可选
COZE_PROJECT_ENV=PROD
DEPLOY_RUN_PORT=5000
COZE_PROJECT_DOMAIN_DEFAULT=https://your-domain.com""")

    add_heading3(doc, '安装与启动')
    add_code_block(doc, """# 安装依赖
pnpm install

# 生产构建
pnpm build

# 启动服务（必须通过 Coze CLI 以获取 AI 认证）
coze dev     # 开发模式
coze start   # 生产模式""")
    add_info_box(doc, '重要', '独立部署时必须通过 Coze CLI（coze dev/coze start）启动服务，才能获取 AI 能力认证。直接 node 运行无法使用 LLM/ASR/TTS/知识库等功能。')

    add_heading3(doc, '数据库初始化')
    add_body(doc, '首次启动时，系统会自动检测数据库状态：', indent=True)
    add_bullet(doc, '如果数据表已存在且管理员账号已创建，无需任何操作，系统正常运行')
    add_bullet(doc, '如果数据表不存在，系统会在控制台输出建表 SQL 语句，需要在 Supabase SQL Editor 中手动执行')
    add_bullet(doc, '数据表创建后，系统会在下次启动时自动创建默认管理员账号（admin / admin123）')
    add_body(doc, '也可以主动查询初始化状态：', indent=True)
    add_code_block(doc, """# 查看数据库初始化状态（返回哪些表已存在、管理员账号是否已创建、建表SQL）
curl http://localhost:5000/api/setup

# 手动触发初始化检查（创建管理员账号等）
curl -X POST http://localhost:5000/api/setup""")
    add_body(doc, '如需手动建表，可在 Supabase SQL Editor 中执行项目中的建表 SQL（6张表 + 默认管理员账号，共约120行），建表后重启服务即可。', indent=True)

    add_heading2(doc, '3.4 验证部署')
    add_body(doc, '无论哪种部署方式，启动后按以下步骤验证：', indent=True)
    add_table(doc,
        ['验证项', '方法', '预期结果'],
        [
            ['服务存活', '浏览器访问部署地址', '显示登录页面'],
            ['数据库连接', '登录管理员后台，查看教师列表', '能正常加载，无报错'],
            ['AI对话', '学生登录后发送文字消息', 'AI能正常回复'],
            ['语音识别', '点击电话按钮说话', 'AI能识别并回应'],
            ['知识库', '教师面板添加文本到知识库', '添加成功，搜索能找到'],
            ['头像上传', '教师面板上传头像照片', '上传成功，头像正常显示'],
        ],
        col_widths=[3, 5, 8]
    )

    # ===== 第四章 管理员操作 =====
    add_heading1(doc, '第四章 管理员操作')

    add_heading2(doc, '4.1 登录管理员后台')
    add_body(doc, '在登录页面点击"管理员入口"，输入用户名（admin）和密码（admin123）登录。首次登录后建议修改默认密码。', indent=True)

    add_heading2(doc, '4.2 教师管理')
    add_bullet(doc, '创建教师账号：点击"添加教师"，填写用户名和密码')
    add_bullet(doc, '编辑教师信息：修改教师账号的用户名或密码')
    add_bullet(doc, '启用/禁用教师：禁用后教师无法登录，但账号数据保留')
    add_bullet(doc, '删除教师：永久删除教师账号及其所有学生数据（知识掌握、对话记录、教学策略）')

    add_heading2(doc, '4.3 学生管理')
    add_bullet(doc, '查看所有学生列表及其关联的助教')
    add_bullet(doc, '启用/禁用学生：禁用后学生无法登录')
    add_bullet(doc, '删除学生：永久删除学生账号及学习记忆数据')

    # ===== 第五章 教师操作 =====
    add_heading1(doc, '第五章 教师操作')

    add_heading2(doc, '5.1 设置向导（首次登录）')
    add_body(doc, '教师首次登录后进入4步设置向导，完成后学生才能看到该助教。', indent=True)

    add_heading3(doc, 'Step 1：基本信息')
    add_bullet(doc, '填写姓名（展示给学生看的名称，如"张老师"）')
    add_bullet(doc, '填写职称（如"高级教师"、"副教授"）')
    add_bullet(doc, '填写教授科目（如"高中物理"、"大学英语"）')

    add_heading3(doc, 'Step 2：授课风格')
    add_bullet(doc, '专业领域：详细描述专业方向（如"力学、电磁学、光学"）')
    add_bullet(doc, '教学风格：描述教学特点（如"善于用生活例子解释抽象概念"）')
    add_bullet(doc, '引导问题：设置学生首次进入时AI助教的开场引导问题')
    add_bullet(doc, '也可选择预设模板：物理教师/数学教师/英语教师/AI计算机教师')

    add_heading3(doc, 'Step 3：上传数字人头像')
    add_bullet(doc, '上传教师真实照片作为数字人形象')
    add_bullet(doc, '语音播报时会在头像周围展示音频波纹动画，直观反馈播放状态')
    add_bullet(doc, '也可使用系统默认头像')

    add_heading3(doc, 'Step 4：初始化知识库')
    add_bullet(doc, '可一键导入5篇示例教学资料（数学/物理/语文/英语/化学）')
    add_bullet(doc, '也可跳过此步，后续在管理面板中添加')

    add_heading2(doc, '5.2 管理面板')
    add_body(doc, '设置完成后进入管理面板，包含3个标签页：', indent=True)

    add_heading3(doc, '助教档案')
    add_bullet(doc, '查看和编辑基本信息、授课风格、引导问题')
    add_bullet(doc, '修改数字人头像')
    add_bullet(doc, '调整声音设置（音色选择、语速、音量）')

    add_heading3(doc, '知识库')
    add_bullet(doc, '文本添加：直接输入教学文本内容')
    add_bullet(doc, '链接导入：输入URL，系统自动抓取网页内容')
    add_bullet(doc, '文件上传：支持 .txt/.md/.csv/.json/.html/.xml/.docx/.pdf 格式，单文件不超过5MB（注意：不支持旧版 .doc 格式，请转换为 .docx）')
    add_bullet(doc, '搜索浏览：输入关键词语义搜索知识库内容')

    add_heading3(doc, '学生记忆')
    add_bullet(doc, '查看所有学生的知识掌握情况（已掌握/学习中/薄弱三档）')
    add_bullet(doc, '查看对话历史记录')
    add_bullet(doc, '了解系统对每位学生的教学策略')

    # ===== 第六章 学生操作 =====
    add_heading1(doc, '第六章 学生操作')

    add_heading2(doc, '6.1 登录与选择助教')
    add_bullet(doc, '输入姓名登录系统')
    add_bullet(doc, '在助教列表中选择教师（展示头像、名称、科目简介）')
    add_bullet(doc, '仅显示已设置完成的助教')

    add_heading2(doc, '6.2 语音交互')
    add_heading3(doc, '开始通话')
    add_bullet(doc, '点击电话按钮启动语音通话')
    add_bullet(doc, '首次使用需授权麦克风权限')
    add_bullet(doc, 'AI助教根据学生记忆主动打招呼（如："上次我们聊到了牛顿第二定律，你还有疑问吗？"）')

    add_heading3(doc, '对话流程')
    add_bullet(doc, '对着麦克风说话，系统自动检测语音起止')
    add_bullet(doc, '说话时页面显示音量指示器')
    add_bullet(doc, 'AI自动识别语音、检索知识库、生成回答并语音播报')
    add_bullet(doc, '播报中：音频波纹动画环绕数字人形象，显示停止按钮')
    add_bullet(doc, '点击停止按钮可手动打断AI播报')

    add_heading3(doc, '交互规则')
    add_bullet(doc, 'AI只回答与教师专业相关的问题，无关问题会被礼貌拒绝')
    add_bullet(doc, '说"懂了"后AI会出验证测试题，答对才标记为已掌握')
    add_bullet(doc, 'AI根据知识掌握情况调整讲解深度（已掌握不重复、薄弱重点讲）')

    add_heading2(doc, '6.3 文字对话')
    add_bullet(doc, '在输入框输入问题，按回车发送')
    add_bullet(doc, 'AI回答以流式文字逐步显示')
    add_bullet(doc, '文字对话同样受价值观领域约束和记忆系统影响')

    add_heading2(doc, '6.4 辅助图示')
    add_body(doc, '当AI判断某个概念适合用图示辅助理解时，会在回答下方显示"生成图示"按钮。点击后系统自动生成辅助图示，帮助学生直观理解抽象概念。', indent=True)

    add_heading2(doc, '6.5 学习记忆')
    add_body(doc, '系统自动记录学生的学习过程，学生可在"学习记忆"面板中查看：', indent=True)
    add_bullet(doc, '知识掌握情况：哪些知识点已掌握、正在学习、比较薄弱')
    add_bullet(doc, '对话历史：与AI助教的历次对话记录')
    add_bullet(doc, '学习画像：系统对学生学习风格和偏好的分析')

    # ===== 第七章 常见问题 =====
    add_heading1(doc, '第七章 常见问题')

    faqs = [
        ('Q: 如何用源码部署项目？', 'A: 最简便的方式是将源码 ZIP 上传到 Coze 编程平台，通过自然语言对话配置数据库和对象存储信息即可，平台自动完成其余配置。详见第三章3.1节。'),
        ('Q: Coze SDK 调用返回认证错误？', 'A: 确保通过 Coze CLI（coze dev）或 Coze 编程平台启动项目。Coze AI 能力的认证由平台自动注入，独立部署直接 node 运行无法获取认证。'),
        ('Q: 启动报错"COZE_SUPABASE_URL is not set"？', 'A: 需要配置 Supabase 数据库连接信息。Coze 编程平台用户可通过自然语言对话配置；独立部署需在项目根目录创建 .env 文件。详见第三章。'),
        ('Q: 头像上传失败？', 'A: 检查 S3 对象存储配置是否正确。Coze 编程平台用户告知平台存储信息即可；独立部署需确认 COZE_BUCKET_ENDPOINT_URL 和 COZE_BUCKET_NAME 已正确配置。'),
        ('Q: 语音通话时AI听不到我说话？', 'A: 请检查浏览器是否授权了麦克风权限。在浏览器地址栏左侧点击锁图标，确认麦克风权限为"允许"，然后刷新页面重试。'),
        ('Q: AI回答了与学习无关的问题？', 'A: 系统已内置价值观领域约束，AI严格限定在教师专业范围内回答。如果仍出现越界回答，教师可在档案中补充专业领域描述，增强约束效果。'),
        ('Q: 每次对话AI都从零开始，不记得我之前学过什么？', 'A: 系统通过学生记忆自动追踪学习进度。请确保每次使用同一姓名登录，系统会自动识别并加载历史记忆。'),
        ('Q: 上传知识库文件失败？', 'A: 请确认文件格式为 .txt/.md/.csv/.json/.html/.xml/.docx/.pdf，且文件大小不超过5MB。注意：不支持旧版 .doc 格式，请将文件另存为 .docx 后重新上传。'),
        ('Q: 学生看不到我创建的助教？', 'A: 请确认已完成4步设置向导。只有设置完成的助教（is_setup_complete=true）才会展示给学生。'),
        ('Q: AI语音播报时我能打断吗？', 'A: 可以。播报期间页面显示带脉冲动画的停止按钮，点击即可打断AI，恢复为倾听状态。'),
        ('Q: 学习记忆数据存储在哪里？', 'A: 学习记忆存储在云端Supabase数据库，按教师隔离。同一学生在不同教师下有独立的学习记忆。'),
        ('Q: 管理员默认账号是什么？', 'A: 用户名: admin，密码: admin123。Coze编程平台部署时自动创建；独立部署时数据表创建后系统自动插入，也可通过 /api/setup 接口查询初始化状态。首次登录后建议立即修改密码。'),
        ('Q: 独立部署后登录管理员提示"用户名或密码错误"？', 'A: 说明数据库表尚未创建。请访问 /api/setup 获取建表SQL，在 Supabase SQL Editor 中执行后重启服务，系统会自动创建默认管理员账号。'),
    ]
    for q, a in faqs:
        add_body(doc, q, bold=True)
        add_body(doc, a, indent=True)

    # 保存
    output = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'assets', '使用手册.docx')
    doc.save(output)
    print(f'使用手册已保存: {output}')
    return output


# ========== 开发手册 ==========

def generate_dev_manual():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('开放智慧助教')
    run.font.name = FONT_TITLE
    run.font.size = Pt(36)
    run.font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_TITLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('开 发 手 册')
    run.font.name = FONT_TITLE
    run.font.size = Pt(30)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_TITLE)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('基于教师个人知识库的智能语音交互教学系统')
    run.font.name = FONT_BODY
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

    doc.add_page_break()

    # ===== 第一章 项目架构 =====
    add_heading1(doc, '第一章 项目架构')

    add_heading2(doc, '1.1 整体架构')
    add_body(doc, '系统采用前后端一体架构（Next.js App Router），前端负责交互与状态管理，后端API路由集成AI能力。整体分为四层：', indent=True)
    add_table(doc,
        ['层次', '技术', '说明'],
        [
            ['用户层', 'React 19 + shadcn/ui', '语音交互界面、管理面板、登录页'],
            ['应用层', 'Next.js 16 API Routes', '对话、记忆、知识库、教师管理等业务逻辑'],
            ['AI能力层', 'Coze SDK', 'LLM对话、ASR语音识别、TTS语音合成、知识库检索、图片生成'],
            ['数据层', 'Supabase + S3', '教师档案、学生记忆、知识库、头像存储'],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading2(doc, '1.2 目录结构')
    add_code_block(doc, """src/
├── app/
│   ├── api/
│   │   ├── audio/asr/route.ts     # 语音识别
│   │   ├── audio/tts/route.ts     # 语音合成
│   │   ├── audio/proxy/route.ts   # 音频代理(CORS)
│   │   ├── chat/route.ts          # 智能对话(核心)
│   │   ├── image/design/route.ts  # 图示设计
│   │   ├── image/generate/route.ts# 图片生成
│   │   ├── image/status/route.ts  # 异步任务状态
│   │   ├── knowledge/add/route.ts # 添加知识库
│   │   ├── knowledge/upload/route.ts # 上传文件
│   │   ├── knowledge/search/route.ts # 搜索知识库
│   │   ├── memory/recall/route.ts # 记忆检索
│   │   ├── memory/update/route.ts # 记忆更新
│   │   ├── memory/profile/route.ts# 学生画像CRUD
│   │   └── teacher/route.ts       # 教师管理
│   ├── page.tsx                   # 主页面(数字教学)
│   ├── layout.tsx                 # 根布局
│   └── globals.css                # 全局样式
├── components/
│   ├── DigitalHuman.tsx           # 数字人(照片+音频波纹动画)
│   ├── VoiceSettings.tsx          # 声音设置
│   ├── KnowledgeManager.tsx       # 知识库管理+学习记忆
│   ├── StudentIdentity.tsx        # 学生身份管理
│   ├── TeacherDashboard.tsx       # 教师管理面板
│   ├── AdminDashboard.tsx         # 管理员后台
│   ├── LoginOverlay.tsx           # 登录页
│   └── ui/                        # shadcn/ui组件库
├── lib/
│   ├── utils.ts                   # 工具函数
│   └── image-task-queue.ts        # 图片异步任务队列
└── storage/database/              # 数据库Schema""")

    add_heading2(doc, '1.3 请求流程')
    add_body(doc, '语音对话的完整请求流程：', indent=True)
    add_code_block(doc, """1. 用户说话 → Web Audio API 采集PCM → WAV编码
2. WAV发送至 /api/audio/asr → 返回识别文字
3. 文字+记忆上下文发送至 /api/chat → LLM流式生成回答
4. 回答按句切分 → 逐句调用 /api/audio/tts → 预解码AudioBuffer
5. AudioBuffer队列顺序播放 → 音频波纹动画反馈
6. 对话结束 → 异步调用 /api/memory/update → LLM提取记忆""")

    # ===== 第二章 技术栈 =====
    add_heading1(doc, '第二章 技术栈与依赖')

    add_heading2(doc, '2.1 核心技术栈')
    add_table(doc,
        ['类别', '技术方案', '版本', '说明'],
        [
            ['前端框架', 'Next.js (App Router)', '16', 'SSR+API路由，前后端一体'],
            ['核心库', 'React', '19', '组件化开发，Hooks状态管理'],
            ['语言', 'TypeScript', '5', '全栈类型安全'],
            ['UI组件', 'shadcn/ui + Radix UI', '-', '无障碍访问，可定制'],
            ['样式', 'Tailwind CSS', '4', '原子化CSS，vintage-grey主题'],
            ['大语言模型', '豆包/DeepSeek', '-', '对话生成，知识检索'],
            ['语音能力', 'Coze ASR/TTS', '-', '语音识别与合成'],
            ['数据库', 'Supabase (PostgreSQL)', '-', '学生记忆持久化'],
            ['对象存储', 'S3兼容存储', '-', '教师头像管理'],
        ],
        col_widths=[3, 5, 2, 6]
    )

    add_heading2(doc, '2.2 关键依赖')
    add_table(doc,
        ['依赖包', '用途'],
        [
            ['coze-coding-dev-sdk', 'AI能力SDK（LLM/ASR/TTS/知识库/图片生成/URL解析）'],
            ['@supabase/supabase-js', 'Supabase客户端，数据库CRUD'],
            ['sonner', 'Toast通知组件'],
            ['docx', 'Word文档生成（报告导出）'],
            ['lucide-react', '图标库'],
        ],
        col_widths=[5, 11]
    )

    # ===== 第三章 核心模块 =====
    add_heading1(doc, '第三章 核心模块详解')

    add_heading2(doc, '3.1 语音交互系统')
    add_body(doc, '语音交互是系统最核心的模块，基于Web Audio API构建持久化音频管道，实现实时语音对话。', indent=True)

    add_heading3(doc, '音频管道架构')
    add_code_block(doc, """麦克风 → ScriptProcessorNode(PCM采集)
       → GainNode(回声防护: speaking时gain=0)
       → AnalyserNode(VAD音量检测 + 音频波纹动画)

TTS播放: AudioBuffer队列 → destination
       → AnalyserNode(音频波纹动画频谱数据)""")

    add_heading3(doc, '关键机制')
    add_table(doc,
        ['机制', '实现方式', '说明'],
        [
            ['回声防护', 'speaking状态gainNode.gain=0', 'TTS播放时静音麦克风，杜绝回声被ASR识别'],
            ['VAD语音检测', '音量阈值+校准+静音超时', '首轮灵敏度优化：noiseFloor初始值3，25分位校准，1.3x阈值'],
            ['TTS预解码', 'fetch→decodeAudioData→队列', '消除段间代理延迟，实现无缝播放'],
            ['短句合并', '<8字暂不发送，等待累积', '减少无效TTS请求，提升语音连贯性'],
            ['手动打断', 'ttsSessionIdRef递增', '点击停止后飞行中TTS自动失效'],
            ['音频波纹动画', 'AnalyserNode频谱数据', 'TTS播报期间实时分析音频频谱，驱动波纹动画'],
            ['有序队列播放', 'ttsSeqRef/ttsNextPlayRef', '确保语音严格按文字顺序播放'],
        ],
        col_widths=[3, 5, 8]
    )

    add_heading2(doc, '3.2 智能对话系统')

    add_heading3(doc, '对话流程')
    add_body(doc, 'Chat API是系统核心入口，接收用户消息后执行以下流程：', indent=True)
    add_code_block(doc, """1. 并行加载: 教师档案 + 学生记忆 + 知识库上下文
2. 拼接系统提示词: 档案+价值观约束+记忆摘要+知识库
3. LLM流式生成回答(SSE协议)
4. 同步返回文字流 → 异步触发TTS
5. 流结束后异步调用 /api/memory/update""")

    add_heading3(doc, '价值观领域约束')
    add_body(doc, '系统提示词中内置严格的领域约束规则，AI助教只回答与教师专业相关的问题：', indent=True)
    add_bullet(doc, '只回答教师专业领域内的问题，拒绝回答无关话题')
    add_bullet(doc, '禁止讨论政治、娱乐、宗教等非教学话题')
    add_bullet(doc, '偏离专业领域时礼貌引导回学习主题')
    add_info_box(doc, '关键', '价值观约束在系统提示词中设定，不受用户输入影响，确保教育安全。')

    add_heading3(doc, '并行化与缓存')
    add_bullet(doc, '教师档案、记忆检索、知识库查询从串行改为并行加载')
    add_bullet(doc, 'LRU缓存：教师档案(5min)、记忆摘要(2min)、知识库检索(3min)、TTS音频(30min)')
    add_bullet(doc, '首次对话延迟减少200~500ms，重复查询零延迟')

    add_heading2(doc, '3.3 学生记忆系统')

    add_heading3(doc, '三阶段工作流')
    add_table(doc,
        ['阶段', '时机', '操作'],
        [
            ['对话前', 'Chat API调用时', '/api/memory/recall → 获取记忆摘要注入系统提示词'],
            ['对话中', 'LLM生成回答时', '根据记忆个性化回复（已掌握不重复、薄弱重点讲）'],
            ['对话后', '流结束后异步', '/api/memory/update → LLM分析对话提取记忆信息'],
        ],
        col_widths=[3, 4, 9]
    )

    add_heading3(doc, '知识掌握三档评价')
    add_table(doc,
        ['档位', 'mastery_level', '教学策略'],
        [
            ['已掌握', '≥ 0.6', '不重复讲解基础概念，可进阶讨论，提供更深层次内容'],
            ['学习中', '0.3 ~ 0.5', '有了解但未完全掌握，需巩固和补充，不需从零开始'],
            ['薄弱', '< 0.3', '理解困难，需重点讲解，放慢节奏，用直观方式说明'],
        ],
        col_widths=[3, 4, 9]
    )

    add_heading3(doc, '掌握度评分机制')
    add_body(doc, '系统通过LLM分析学生对问题的回复，自动评价掌握程度并更新mastery_level数值：', indent=True)
    add_table(doc,
        ['评价类型', '触发条件', 'mastery_level 变化', '说明'],
        [
            ['mastered', '学生通过验证测试（答对问题）', '已有记录设为0.8；新建时设为0.7', '必须通过测试验证，不能仅凭学生说懂了'],
            ['practice', '学生说懂了但未测试，或一般性练习', '+0.1（上限0.6，不超过已掌握阈值）', '有进步但未经验证，不超过已掌握档位'],
            ['confused', '学生仍然困惑或答错', '-0.1（下限0）', '降低掌握度，下次重点讲解'],
        ],
        col_widths=[3, 4, 4, 5]
    )

    add_info_box(doc, '验证测试机制', '学生说"懂了"时，AI必须出验证测试题，通过后才标记为已掌握。这是防止学生虚报掌握程度的关键机制。')

    add_heading3(doc, '重复记录合并')
    add_body(doc, 'LLM分析对话提取知识点时，可能产生重复记录。系统采用三级模糊匹配自动合并：', indent=True)
    add_bullet(doc, '第一级：精确匹配（topic完全相同）')
    add_bullet(doc, '第二级：subtopic关键词匹配（关键词重叠）')
    add_bullet(doc, '第三级：topic+subtopic全文关键词匹配')
    add_bullet(doc, '合并策略：保留最佳记录，合并practice_count，删除重复记录')

    add_heading3(doc, '数据隔离')
    add_body(doc, '所有记忆数据按teacher_id隔离：同一学生在不同教师下有独立的画像、知识掌握、对话记录和教学策略。', indent=True)

    add_heading2(doc, '3.4 图片生成')
    add_body(doc, '图片生成采用异步任务队列模式，避免前端长时间等待：', indent=True)
    add_code_block(doc, """1. 前端提交 → /api/image/generate (async:true)
2. 立即返回 task_id
3. 后台执行: /api/image/design (LLM生成prompt)
           → Coze Image Generation (生成图片)
4. 前端轮询 /api/image/status?task_id=xxx → 获取结果""")

    # ===== 第四章 API接口 =====
    add_heading1(doc, '第四章 API接口文档')

    add_heading2(doc, '4.1 接口总览')
    add_table(doc,
        ['接口', '方法', '路径', '功能'],
        [
            ['语音识别', 'POST', '/api/audio/asr', '语音转文字(WAV)'],
            ['语音合成', 'POST', '/api/audio/tts', '文字转语音(多音色)'],
            ['音频代理', 'GET', '/api/audio/proxy', '代理TTS URL(CORS)'],
            ['智能对话', 'POST', '/api/chat', '知识库问答(流式)'],
            ['图示设计', 'POST', '/api/image/design', 'LLM生成图示prompt'],
            ['图片生成', 'POST', '/api/image/generate', '图片生成(同步/异步)'],
            ['任务状态', 'GET', '/api/image/status', '异步任务状态查询'],
            ['添加文档', 'POST', '/api/knowledge/add', '文本/URL添加知识库'],
            ['上传文件', 'POST', '/api/knowledge/upload', '文件上传(.txt/.md/.csv/.json/.html/.xml/.docx/.pdf, FormData≤5MB)'],
            ['搜索文档', 'POST', '/api/knowledge/search', '语义搜索知识库'],
            ['记忆检索', 'POST', '/api/memory/recall', '对话前加载记忆上下文'],
            ['记忆更新', 'POST', '/api/memory/update', 'LLM分析对话后写入记忆'],
            ['学生画像', 'GET/POST/DELETE', '/api/memory/profile', '学生画像CRUD'],
            ['教师管理', 'GET/POST', '/api/teacher', '教师登录/档案/学生列表'],
            ['教师头像', 'GET/POST', '/api/teacher/avatar', '头像上传/签名URL'],
            ['管理员', 'GET/POST', '/api/admin', '管理员登录/教师学生管理'],
        ],
        col_widths=[3, 3, 5, 5]
    )

    add_heading2(doc, '4.2 核心API详解')

    add_heading3(doc, 'POST /api/chat')
    add_body(doc, '智能对话核心接口，基于知识库和记忆系统生成个性化回答。', indent=True)
    add_bullet(doc, '请求参数：message(用户消息)、teacher_id(教师ID)、student_id(学生ID)、session_id(会话ID)')
    add_bullet(doc, '响应格式：SSE流式输出(text/event-stream)')
    add_bullet(doc, '内部流程：并行加载→拼接提示词→LLM流式生成→异步记忆更新')
    add_bullet(doc, '价值观约束：系统提示词严格限定AI只回答专业问题')

    add_heading3(doc, 'POST /api/memory/recall')
    add_body(doc, '对话前检索学生记忆，生成三档分类摘要注入系统提示词。', indent=True)
    add_bullet(doc, '请求参数：student_id、teacher_id')
    add_bullet(doc, '返回：学生画像、知识掌握（已掌握/学习中/薄弱三档）、近期对话摘要、教学策略')
    add_bullet(doc, '数据隔离：按teacher_id过滤，同一学生在不同教师下有独立记忆')

    add_heading3(doc, 'POST /api/memory/update')
    add_body(doc, '对话结束后由LLM自动分析对话内容，提取记忆信息写入数据库。', indent=True)
    add_bullet(doc, '请求参数：student_id、teacher_id、对话内容')
    add_bullet(doc, 'LLM分析维度：知识点掌握度(mastered/practice/confused)、教学策略、学生画像更新')
    add_bullet(doc, '自动合并：三级模糊匹配检测重复知识点，合并practice_count后删除重复记录')

    # ===== 第五章 数据库 =====
    add_heading1(doc, '第五章 数据库设计')

    add_heading2(doc, '5.1 核心表')
    add_table(doc,
        ['表名', '说明', '关键字段'],
        [
            ['teacher_profile', '教师档案', 'id, name, password, subjects, expertise, teaching_style, avatar_key, knowledge_table, is_setup_complete, is_enabled'],
            ['student_profile', '学生画像', 'id, student_id, teacher_id, name, learning_style, interests, goals'],
            ['knowledge_mastery', '知识掌握追踪', 'id, student_id, teacher_id, topic, subtopic, mastery_level(0~1), practice_count, strong_points, weak_points'],
            ['conversation_log', '对话记录', 'id, student_id, teacher_id, session_id, role, content, timestamp'],
            ['teaching_strategy', '教学策略记忆', 'id, student_id, teacher_id, method_name, description, effectiveness'],
            ['admin_account', '管理员账号', 'id, username, password, created_at'],
        ],
        col_widths=[4, 3, 9]
    )

    add_heading2(doc, '5.2 知识库隔离')
    add_body(doc, '每位教师有独立的知识库表（表名=教师ID），知识添加/搜索/对话都使用各自专属表，确保教师间数据完全隔离。', indent=True)

    add_heading2(doc, '5.3 知识掌握评价核心字段')
    add_table(doc,
        ['字段', '类型', '说明'],
        [
            ['mastery_level', 'FLOAT(0~1)', '掌握度数值：≥0.6已掌握, 0.3~0.5学习中, <0.3薄弱'],
            ['practice_count', 'INTEGER', '练习次数，合并重复记录时累加'],
            ['strong_points', 'TEXT', '已掌握项的优势要点（验证测试通过后设置）'],
            ['weak_points', 'TEXT', '薄弱项/学习中的薄弱要点'],
        ],
        col_widths=[4, 3, 9]
    )

    # ===== 第六章 AI集成 =====
    add_heading1(doc, '第六章 AI集成规范')

    add_heading2(doc, '6.1 SDK使用规范')
    add_bullet(doc, 'coze-coding-dev-sdk 必须在后端API路由中使用，禁止客户端直接调用')
    add_bullet(doc, '所有API调用使用 HeaderUtils.extractForwardHeaders 转发请求头')
    add_bullet(doc, '对话API默认SSE流式输出（text/event-stream）')
    add_bullet(doc, '所有API调用包含try-catch错误处理')

    add_heading2(doc, '6.2 系统提示词结构')
    add_code_block(doc, """系统提示词组成:
1. 教师档案（姓名、科目、专业、教学风格）
2. 价值观领域约束（只回答专业问题，拒绝无关话题）
3. 记忆摘要（已掌握/学习中/薄弱三档，附交互规则）
4. 知识库检索结果（RAG相关文档片段）
5. 交互规则（不重复已掌握、重点讲薄弱、验证测试等）""")

    add_heading2(doc, '6.3 RAG架构')
    add_body(doc, '知识库采用检索增强生成（RAG）架构：', indent=True)
    add_bullet(doc, '教师上传文本/文件/URL → Coze Embedding向量化 → 存入教师专属知识库表')
    add_bullet(doc, '对话时自动检索相关知识库内容 → 注入系统提示词 → LLM结合知识生成回答')
    add_bullet(doc, '搜索API支持语义搜索，兼容旧表名teacher_knowledge_base')

    add_heading2(doc, '6.4 ASR/TTS集成')
    add_bullet(doc, 'ASR：接收WAV音频，调用Coze ASR转换文字，区分"无语音"和"格式错误"')
    add_bullet(doc, 'TTS：调用Coze TTS合成语音，speaker白名单验证，失败自动回退默认音色')
    add_bullet(doc, '音频代理：/api/audio/proxy解决CORS问题，支持TTS URL代理访问')

    add_heading2(doc, '6.5 LRU缓存配置')
    add_table(doc,
        ['缓存对象', 'TTL', '说明'],
        [
            ['教师档案', '5分钟', '避免每次对话查数据库'],
            ['记忆摘要', '2分钟', '同一学生短时间内复用'],
            ['知识库检索', '3分钟', '相似问题复用检索结果'],
            ['TTS音频', '30分钟', '相同文本+音色复用语音'],
        ],
        col_widths=[4, 3, 9]
    )

    # ===== 第七章 部署 =====
    add_heading1(doc, '第七章 部署与运维')

    add_heading2(doc, '7.1 环境变量')
    add_table(doc,
        ['变量名', '说明', '示例'],
        [
            ['COZE_WORKSPACE_PATH', '项目工作目录', '/workspace/projects/'],
            ['COZE_PROJECT_DOMAIN_DEFAULT', '对外访问域名', 'https://xxx.dev.coze.site'],
            ['DEPLOY_RUN_PORT', '服务监听端口', '5000'],
            ['COZE_PROJECT_ENV', '运行环境', 'DEV / PROD'],
        ],
        col_widths=[5, 4, 7]
    )

    add_heading2(doc, '7.2 构建与启动')
    add_code_block(doc, """# 开发环境
coze dev              # 启动开发服务器(端口5000, HMR热更新)

# 生产环境
coze build            # 构建生产版本
coze start            # 启动生产服务器

# 包管理(仅pnpm)
pnpm install          # 安装依赖
pnpm add <package>    # 添加依赖""")

    add_heading2(doc, '7.3 数据库初始化')
    add_body(doc, '项目首次部署时，Supabase数据库表和默认管理员账号会自动创建。默认管理员账号：admin / admin123。', indent=True)

    add_heading2(doc, '7.4 开发规范')
    add_heading3(doc, 'Ref同步模式')
    add_body(doc, '为避免useCallback闭包捕获旧状态值，所有跨回调共享的状态创建对应ref，渲染阶段同步赋值，回调内部通过ref.current读取最新值。', indent=True)

    add_heading3(doc, 'Hydration问题防范')
    add_bullet(doc, '严禁JSX中直接使用typeof window、Date.now()、Math.random()')
    add_bullet(doc, '必须使用use client + useEffect + useState确保客户端渲染')
    add_bullet(doc, '禁止非法HTML嵌套（如<p>嵌套<div>）')

    add_heading3(doc, '包管理约束')
    add_body(doc, '仅允许使用pnpm作为包管理器，严禁使用npm或yarn。', indent=True)

    # 保存
    output = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'assets', '开发手册.docx')
    doc.save(output)
    print(f'开发手册已保存: {output}')
    return output


# ========== 主入口 ==========
if __name__ == '__main__':
    generate_user_manual()
    generate_dev_manual()
    print('\n全部完成!')
